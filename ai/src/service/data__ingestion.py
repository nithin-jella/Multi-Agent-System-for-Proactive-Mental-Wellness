import re
import io
import logging
from typing_extensions import Literal
import re

import pymupdf

import docx
from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import zipfile
import xml.etree.ElementTree as ET

from html.parser import HTMLParser
from html import unescape

logger = logging.getLogger(__name__)

class DataIngestion:
    def __init__(self):
        pass

    def extract_text_from_files(self, file_path: str, type: Literal["pdf", "docx", "html"], split: bool = True) -> list[str]:
        logger.info("Start Text Extraction")
        if type == "pdf":
            return self.__extract_text_from_pdf(file_path=file_path, split = split)
        if type == "docx":
            return self.__extract_text_from_docx(docx_path=file_path, split = split)
        if type == "html":
            return self.__extract_text_from_html(file_path=file_path, split = split)
        
    @staticmethod
    def __extract_text_from_html(file_path: str, split: bool = True) -> list[str]:
        """
        Extract clean text from HTML content, removing all tags and styling.
        
        Args:
            file_path (str): HTML string to extract text from
            
        Returns:
            tuple: (extracted_text, sections_list)
        """
        
        class HTMLTextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self.text_parts = []
                self.block_elements = {
                    'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div', 
                    'article', 'section', 'header', 'footer', 'main',
                    'blockquote', 'pre', 'ul', 'ol', 'li', 'dl', 'dt', 'dd',
                    'table', 'tr', 'td', 'th', 'tbody', 'thead', 'tfoot'
                }
                self.ignore_tags = {
                    'style', 'script', 'meta', 'link', 'noscript', 'head', 'title',
                    'svg', 'canvas', 'iframe', 'embed', 'object'
                }
                self.list_tags = {'ul', 'ol'}
                self.list_item_tags = {'li'}
                
                # State tracking
                self.current_tag = None
                self.ignore_content = False
                
                # Table handling
                self.in_table = False
                self.in_cell = False
                self.current_row = []
                self.current_table = []
                self.current_cell_content = []
                
                # List handling
                self.in_list = False
                self.list_stack = []  # Track nested lists
                self.list_item_counter = 0
                
                # Headings for sections
                self.heading_levels = {'h1': 1, 'h2': 2, 'h3': 3, 'h4': 4, 'h5': 5, 'h6': 6}
                
            def clean_text(self, text):
                """Clean and normalize text content"""
                if not text:
                    return ""
                
                # Decode HTML entities
                text = unescape(text)
                
                # Remove problematic characters
                text = text.replace('\xa0', ' ')     # Non-breaking space
                text = text.replace('\u00a0', ' ')   # Another non-breaking space
                text = text.replace('\u200b', '')    # Zero-width space
                text = text.replace('\ufeff', '')    # Byte order mark
                text = text.replace('�', '')         # Replacement character
                text = text.replace('Â', '')         # Malformed UTF-8
                text = text.replace('â€™', "'")      # Malformed apostrophe
                text = text.replace('â€œ', '"')      # Malformed opening quote
                text = text.replace('â€\x9d', '"')   # Malformed closing quote
                text = text.replace('\r\n', '\n')    # Normalize line endings
                text = text.replace('\r', '\n')      # Normalize line endings
                
                # Clean whitespace but preserve meaningful spaces
                text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space
                text = re.sub(r'\n\s*\n', '\n\n', text)  # Clean up multiple newlines
                text = text.strip()
                
                return f"{text} "
                
            def handle_starttag(self, tag, attrs):
                self.current_tag = tag.lower()
                
                # Check if we should ignore content inside this tag
                if self.current_tag in self.ignore_tags:
                    self.ignore_content = True
                    return
                
                # Handle headings
                if self.current_tag in self.heading_levels:
                    # Add H1 marker for section splitting (backward compatibility)
                    if self.current_tag == 'h1':
                        self.text_parts.append('H1')
                    # Add line break before heading
                    if self.text_parts and not self.text_parts[-1].endswith('\n'):
                        self.text_parts.append('\n')
                
                # Handle lists
                if self.current_tag in self.list_tags:
                    self.in_list = True
                    self.list_stack.append(self.current_tag)
                    self.list_item_counter = 0
                    if self.text_parts and not self.text_parts[-1].endswith('\n'):
                        self.text_parts.append('\n')
                
                if self.current_tag in self.list_item_tags and self.in_list:
                    self.list_item_counter += 1
                    # Add appropriate list marker
                    if self.list_stack and self.list_stack[-1] == 'ol':
                        self.text_parts.append(f"{self.list_item_counter}. ")
                    else:
                        self.text_parts.append("• ")
                
                # Handle tables
                if self.current_tag == 'table':
                    self.in_table = True
                    self.current_table = []
                    # Add spacing before table
                    if self.text_parts and not self.text_parts[-1].endswith('\n'):
                        self.text_parts.append('\n')

                if self.current_tag == 'tr' and self.in_table:
                    self.current_row = []

                if self.current_tag in {'td', 'th'} and self.in_table:
                    self.in_cell = True
                    self.current_cell_content = []

                # Add line break before block elements (except list items handled above)
                if (self.current_tag in self.block_elements and 
                    self.current_tag not in self.list_item_tags and
                    self.current_tag not in {'table', 'tr', 'td', 'th'}):
                    if self.text_parts and not self.text_parts[-1].endswith('\n'):
                        self.text_parts.append('\n')
                        
            def handle_endtag(self, tag):
                tag_lower = tag.lower()
                
                # Handle table cells
                if tag_lower in {'td', 'th'} and self.in_cell:
                    self.in_cell = False
                    # Join all content collected for this cell
                    cell_text = ' '.join(self.current_cell_content).strip()
                    self.current_row.append(cell_text)
                    self.current_cell_content = []

                # Handle table rows
                if tag_lower == 'tr' and self.in_table:
                    if self.current_row:  # Only add non-empty rows
                        self.current_table.append(self.current_row)
                    self.current_row = []

                # Handle table end
                if tag_lower == 'table':
                    self.in_table = False
                    if self.current_table and any(row for row in self.current_table if any(cell.strip() for cell in row)):
                        table_text = self._format_table_as_markdown(self.current_table)
                        self.text_parts.append(table_text)
                    # Reset table state
                    self.current_table = [] 
                    self.current_row = []
                
                # Handle list end
                if tag_lower in self.list_tags and self.in_list:
                    if self.list_stack:
                        self.list_stack.pop()
                    if not self.list_stack:
                        self.in_list = False
                    self.text_parts.append('\n')
                
                if tag_lower in self.list_item_tags and self.in_list:
                    self.text_parts.append('\n')
                
                # Handle headings
                if tag_lower in self.heading_levels:
                    self.text_parts.append('\n')
                    
                # Stop ignoring content when we close an ignored tag
                if tag_lower in self.ignore_tags:
                    self.ignore_content = False
                    self.current_tag = None
                    return
                    
                # Add line break after block elements
                if (tag_lower in self.block_elements and 
                    tag_lower not in {'table', 'tr', 'td', 'th'} and
                    tag_lower not in self.list_item_tags):
                    if self.text_parts and not self.text_parts[-1].endswith('\n'):
                        self.text_parts.append('\n')
                        
                self.current_tag = None
            
            def _format_table_as_markdown(self, table_rows):
                """Format table rows as markdown table"""
                if not table_rows:
                    return ""
                
                max_cols = max(len(row) for row in table_rows)
                col_widths = [0] * max_cols
                
                for row in table_rows:
                    while len(row) < max_cols:
                        row.append("")
                
                for row in table_rows:
                    for i, cell in enumerate(row):
                        if i < len(col_widths):
                            col_widths[i] = max(col_widths[i], len(str(cell)))
                
                table_lines = ['\n']
                
                for i, row in enumerate(table_rows):
                    # Format cells with padding
                    formatted_cells = []
                    for j, cell in enumerate(row):
                        if j < len(col_widths):
                            formatted_cells.append(str(cell).ljust(col_widths[j]))
                        else:
                            formatted_cells.append(str(cell))
                    
                    line = "| " + " | ".join(formatted_cells) + " |"
                    table_lines.append(line)
                    
                    if i == 0 and len(table_rows) > 1:
                        separator_cells = []
                        for width in col_widths:
                            separator_cells.append("-" * max(3, width))
                        separator = "| " + " | ".join(separator_cells) + " |"
                        table_lines.append(separator)
                
                table_lines.append('\n')
                return '\n'.join(table_lines)
                
            def handle_data(self, data):
                if self.ignore_content:
                    return
                
                cleaned_data = self.clean_text(data)
                
                if not cleaned_data:
                    return

                if self.in_table and self.in_cell:
                    self.current_cell_content.append(cleaned_data)
                else:
                    self.text_parts.append(cleaned_data)
                    
            def get_text(self):
                """Get complete extracted text"""
                text = ''.join(self.text_parts)
                
                text = self.clean_text(text)
                
                text = re.sub(r'\n{3,}', '\n\n', text)
                
                return text
            
            def get_sections(self):
                """Get sections split by H1 headings"""
                sections = []
                current_section = []
                
                for part in self.text_parts:
                    if part == "H1":
                        if current_section:
                            section_text = ''.join(current_section)
                            section_text = self.clean_text(section_text)
                            if section_text:
                                sections.append(section_text)
                        current_section = []
                    else:
                        current_section.append(part)
                
                if current_section:
                    section_text = ''.join(current_section)
                    section_text = self.clean_text(section_text)
                    if section_text:
                        sections.append(section_text)
                
                return sections
        
        if not file_path or not file_path.strip():
            return "", []
        
        try:
            parser = HTMLTextExtractor()
            f =  open(f"./data/{file_path}", "r", encoding='utf-8')
            parser.feed(f.read())

            result = []
            if (not split):
                result  = [parser.get_text()]
            else:
                result  = parser.get_sections()
            
            return result
            
        except Exception as e:
            logger.error(f"Fail to extract HTML file {e}")
            return []
        finally:
            f.close()
            parser.close()


    @staticmethod
    def __extract_text_from_pdf(file_path: str, split: bool = True) -> list[str] :
        """Extract raw text from  PDF file return text per page as list"""

        try:
            text = []
            file_byte =  open(f"./data/{file_path}", "rb")
            with pymupdf.open(stream=file_byte.read(), filetype="pdf") as doc:
                for page in doc:
                    text.append(page.get_text())
            
            if (not split):
                return "\n\n".join(text)
            return text
        except Exception as e:
            logger.error(f"Failed to extract PDF: {e}")
            raise
        finally:
            file_byte.close()

    @staticmethod
    def __extract_text_from_docx(docx_path: str, split: bool = True) -> list[str]:
        """
        Extract clean text from DOCX content, handling paragraphs, tables, and headers.
        
        Args:
            docx_path (str): Path to DOCX file or file-like object
            
        Returns:
            tuple: (extracted_text, sections_list)
        """
        
        class DocxTextExtractor:
            def __init__(self):
                self.text_parts = []
                self.sections = []
                self.current_section = []
                
            def clean_text(self, text):
                """Clean and normalize text content"""
                if not text:
                    return ""
                    
                # Remove problematic characters
                text = text.replace('\xa0', ' ')  # Non-breaking space
                text = text.replace('\u00a0', ' ')  # Another non-breaking space
                text = text.replace('\u200b', '')  # Zero-width space
                text = text.replace('\ufeff', '')  # Byte order mark
                text = text.replace('�', '')
                text = text.replace('\r', '\n')  # Convert carriage returns
                
                # Clean whitespace but preserve line breaks
                text = re.sub(r'[ \t]+', ' ', text)  # Replace multiple spaces/tabs with single space
                text = re.sub(r'\n\s*\n', '\n\n', text)  # Clean up multiple newlines
                text = text.strip()
                
                return text
            
            def extract_table_text(self, table):
                """Extract text from a table and format as markdown"""
                table_rows = []
                
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        # Extract text from each cell, handling merged cells
                        cell_text = ""
                        for paragraph in cell.paragraphs:
                            para_text = paragraph.text.strip()
                            if para_text:
                                cell_text += para_text + " "
                        
                        # Clean the cell text
                        cell_text = self.clean_text(cell_text)
                        row_cells.append(cell_text)
                    
                    if any(cell.strip() for cell in row_cells):  # Only add non-empty rows
                        table_rows.append(row_cells)
                
                if not table_rows:
                    return ""
                
                # Format as markdown table
                if table_rows:
                    # Calculate column widths
                    max_cols = max(len(row) for row in table_rows) if table_rows else 0
                    col_widths = [0] * max_cols
                    
                    # Pad rows to have same number of columns
                    for row in table_rows:
                        while len(row) < max_cols:
                            row.append("")
                    
                    # Calculate maximum width for each column
                    for row in table_rows:
                        for i, cell in enumerate(row):
                            if i < len(col_widths):
                                col_widths[i] = max(col_widths[i], len(str(cell)))
                    
                    # Format table
                    table_text = "\n"
                    for i, row in enumerate(table_rows):
                        # Format cells with padding
                        formatted_cells = []
                        for j, cell in enumerate(row):
                            if j < len(col_widths):
                                formatted_cells.append(str(cell).ljust(col_widths[j]))
                            else:
                                formatted_cells.append(str(cell))
                        
                        line = "| " + " | ".join(formatted_cells) + " |"
                        table_text += line + "\n"
                        
                        # Add header separator after first row
                        if i == 0:
                            separator_cells = []
                            for width in col_widths:
                                separator_cells.append("-" * max(3, width))
                            separator = "| " + " | ".join(separator_cells) + " |"
                            table_text += separator + "\n"
                    
                    table_text += "\n"
                    return table_text
                
                return ""
            
            def is_heading(self, paragraph):
                """Check if paragraph is a heading"""
                if paragraph.style and paragraph.style.name:
                    style_name = paragraph.style.name.lower()
                    return (style_name.startswith('heading') or 
                        style_name.startswith('title') or
                        'heading' in style_name)
                return False
            
            def get_heading_level(self, paragraph):
                """Get heading level (1-6)"""
                if paragraph.style and paragraph.style.name:
                    style_name = paragraph.style.name.lower()
                    # Try to extract number from heading style
                    import re
                    match = re.search(r'heading\s*(\d+)', style_name)
                    if match:
                        return min(int(match.group(1)), 6)
                    elif 'title' in style_name:
                        return 1
                return 1
            
            def extract_from_document(self, doc: docx.Document):
                """Extract text from document maintaining structure"""
                
                # Iterate through all elements in document order
                for element in doc.element.body:
                    if isinstance(element, CT_P):  # Paragraph
                        paragraph = Paragraph(element, doc)
                        text = paragraph.text.strip()
                        
                        if text:
                            cleaned_text = self.clean_text(text)
                            
                            if self.is_heading(paragraph):
                                # Handle headings
                                level = self.get_heading_level(paragraph)
                                
                                # Save current section if it has content
                                if self.current_section:
                                    section_text = "\n".join(self.current_section)
                                    if section_text.strip():
                                        self.sections.append(section_text.strip())
                                    self.current_section = []
                                
                                # Add heading marker for H1 (for compatibility with your existing code)
                                if level == 1:
                                    self.text_parts.append("H1")
                                
                                # Format heading with markdown
                                heading_text = "#" * level + " " + cleaned_text
                                self.text_parts.append(heading_text)
                                self.current_section.append(heading_text)
                            else:
                                # Regular paragraph
                                self.text_parts.append(cleaned_text)
                                self.current_section.append(cleaned_text)
                    
                    elif isinstance(element, CT_Tbl):  # Table
                        table = Table(element, doc)
                        table_text = self.extract_table_text(table)
                        if table_text.strip():
                            self.text_parts.append(table_text)
                            self.current_section.append(table_text)
                
                # Don't forget the last section
                if self.current_section:
                    section_text = "\n".join(self.current_section)
                    if section_text.strip():
                        self.sections.append(section_text.strip())
            
            def get_text(self):
                """Get complete extracted text"""
                # Join all text parts
                text = "\n".join(part for part in self.text_parts if part.strip())
                return self.clean_text(text)
            
            def get_sections(self):
                """Get sections split by headings"""
                # Clean sections
                cleaned_sections = []
                for section in self.sections:
                    cleaned = self.clean_text(section)
                    if cleaned:
                        cleaned_sections.append(cleaned)
                return cleaned_sections
        
        try:
            print("Extraxt DOCX file")
            # Try to open as DOCX file
            if isinstance(docx_path, str):
                doc = Document(f"./data/{docx_path}")
            else:
                # Assume it's a file-like object
                doc = Document(f"./data/{docx_path}")
            
            extractor = DocxTextExtractor()
            extractor.extract_from_document(doc)

            result = []

            if (not split):
                result = [extractor.get_text()]
            else:
                result = extractor.get_sections()
            return result
        except Exception as e:
            logger.error("Fail to extract DOCX file")
            return []